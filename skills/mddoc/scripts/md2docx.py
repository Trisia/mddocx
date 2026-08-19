#!/usr/bin/env python3
"""
mddoc - Markdown 转学术格式 DOCX

用法:
    python md2docx.py <input.md>          # 转换 Markdown 文件
    python md2docx.py <input.md> -o out.docx  # 指定输出路径
    python md2docx.py --text "markdown..."  # 直接转换文本

依赖:
    pip install python-docx Pillow requests mistune
"""

import argparse
import os
import re
import sys
import tempfile
import urllib.parse
from io import BytesIO

import mistune
from mistune.plugins import table as mistune_table, math as mistune_math

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree
from PIL import Image


# ============================================================
# OMML 公式工具 (LaTeX → OMML XML)
# ============================================================

# OMML 命名空间
NSM = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

def _m_qn(local):
    """构造 m 命名空间的 Clark 格式属性名，如 _m_qn('val') → '{NSM}val'"""
    return f'{{{NSM}}}{local}'

# LaTeX 希腊字母 → Unicode 映射
GREEK = {
    'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'delta': 'δ',
    'epsilon': 'ε', 'zeta': 'ζ', 'eta': 'η', 'theta': 'θ',
    'iota': 'ι', 'kappa': 'κ', 'lambda': 'λ', 'mu': 'μ',
    'nu': 'ν', 'xi': 'ξ', 'omicron': 'ο', 'pi': 'π',
    'rho': 'ρ', 'sigma': 'σ', 'tau': 'τ', 'upsilon': 'υ',
    'phi': 'φ', 'chi': 'χ', 'psi': 'ψ', 'omega': 'ω',
    'Gamma': 'Γ', 'Delta': 'Δ', 'Theta': 'Θ', 'Lambda': 'Λ',
    'Xi': 'Ξ', 'Pi': 'Π', 'Sigma': 'Σ', 'Upsilon': 'Υ',
    'Phi': 'Φ', 'Psi': 'Ψ', 'Omega': 'Ω',
    'varepsilon': 'ϵ', 'varphi': 'ϕ', 'vartheta': 'ϑ',
    'varpi': 'ϖ', 'varrho': 'ϱ', 'varsigma': 'ς',
}

# LaTeX 符号 → Unicode
SYMBOLS = {
    'infty': '∞', 'cdot': '·', 'times': '×', 'pm': '±', 'mp': '∓',
    'leq': '≤', 'geq': '≥', 'neq': '≠', 'approx': '≈', 'equiv': '≡',
    'propto': '∝', 'sim': '∼', 'simeq': '≃',
    'partial': '∂', 'nabla': '∇', 'forall': '∀', 'exists': '∃',
    'in': '∈', 'notin': '∉', 'subset': '⊂', 'supset': '⊃',
    'subseteq': '⊆', 'supseteq': '⊇', 'cap': '∩', 'cup': '∪',
    'rightarrow': '→', 'leftarrow': '←', 'Rightarrow': '⇒', 'Leftarrow': '⇐',
    'leftrightarrow': '↔', 'mapsto': '↦',
    'ldots': '…', 'cdots': '⋯', 'vdots': '⋮', 'ddots': '⋱',
    'circ': '∘', 'bullet': '•', 'oplus': '⊕', 'otimes': '⊗',
    'angle': '∠', 'triangle': '△', 'square': '□',
    'mid': '|', 'parallel': '∥', 'perp': '⊥',
    'aleph': 'ℵ', 'hbar': 'ℏ', 'imath': 'ı', 'jmath': 'ȷ',
    'ell': 'ℓ', 'wp': '℘', 'Re': 'ℜ', 'Im': 'ℑ',
    'prime': '′', 'emptyset': '∅', 'varnothing': '∅',
    'langle': '⟨', 'rangle': '⟩', 'lceil': '⌈', 'rceil': '⌉',
    'lfloor': '⌊', 'rfloor': '⌋',
    # 阶段3 补全：二元运算符
    'div': '÷', 'ominus': '⊖', 'oslash': '⊘', 'star': '⋆', 'odot': '⊙',
    # 阶段3 补全：关系符
    'll': '≪', 'gg': '≫', 'prec': '≺', 'succ': '≻',
    'preceq': '≼', 'succeq': '≽', 'nsim': '≁', 'asymp': '≍',
    # 阶段3 补全：逻辑符号
    'top': '⊤', 'bot': '⊥', 'neg': '¬', 'nexists': '∄',
    'wedge': '∧', 'vee': '∨',
    # 阶段3 补全：集合符号
    'setminus': '∖', 'nsubseteq': '⊈', 'nsupseteq': '⊉',
    # 阶段3 补全：箭头
    'to': '→', 'longrightarrow': '⟶', 'longmapsto': '⟼',
    'uparrow': '↑', 'downarrow': '↓', 'updownarrow': '↕',
    'Uparrow': '⇑', 'Downarrow': '⇓', 'Updownarrow': '⇕',
    'longleftarrow': '⟵', 'Leftrightarrow': '⇔',
    # 阶段3 补全：常用别名
    'ge': '≥', 'le': '≤', 'ne': '≠', 'gets': '←',
    'backslash': '\\',
}

# 需要特殊处理的函数名
FUNCTIONS = {
    'sin', 'cos', 'tan', 'cot', 'sec', 'csc',
    'arcsin', 'arccos', 'arctan',
    'sinh', 'cosh', 'tanh', 'coth',
    'log', 'lg', 'ln', 'exp',
    'lim', 'max', 'min', 'sup', 'inf',
    'det', 'dim', 'ker', 'gcd', 'deg', 'hom',
    'arg', 'Pr',
}

BIG_OPS = {
    'sum': '∑', 'prod': '∏', 'coprod': '∐',
    'int': '∫', 'iint': '∬', 'iiint': '∭', 'oint': '∮',
    'bigcup': '⋃', 'bigcap': '⋂', 'bigvee': '⋁', 'bigwedge': '⋀',
    'bigoplus': '⨁', 'bigotimes': '⨂', 'bigodot': '⨀',
}

ACCENTS = {
    'hat': '̂', 'bar': '̅', 'vec': '⃗', 'dot': '̇',
    'ddot': '̈', 'tilde': '̃', 'breve': '̆', 'check': '̌',
    'acute': '́', 'grave': '̀',
}


def _m_elem(tag):
    """创建带 OMML 命名空间的 XML 元素"""
    return etree.Element(f'{{{NSM}}}{tag}', nsmap={'m': NSM})


def _m_run(text, italic=True):
    """创建 m:r 元素，包含 m:t 文本子元素"""
    r = _m_elem('r')
    if not italic:
        rPr = _m_elem('rPr')
        nor = _m_elem('nor')
        nor.set(_m_qn('val'), '1')
        rPr.append(nor)
        r.append(rPr)
    t_elem = _m_elem('t')
    t_elem.set(qn('xml:space'), 'preserve')
    t_elem.text = text
    r.append(t_elem)
    return r


def _m_text_run(text):
    """创建非斜体文本 run（用于 \text{}、函数名等）"""
    return _m_run(text, italic=False)


def _m_wrap_d(children):
    """将子元素列表包装在 m:d（分隔符）元素中"""
    d = _m_elem('d')
    for c in children:
        if isinstance(c, str):
            d.append(_m_run(c))
        elif c is not None:
            d.append(c)
    return d


def _m_run_with_style(text, italic=True, bold=False, script=None, normal=False):
    """创建带样式属性的 m:r 元素

    参数:
        text: 显示文本
        italic: 是否斜体（默认 True）
        bold: 是否粗体
        script: 脚本字体类型 ('script' 对应 \\mathcal, 'double-struck' 对应 \\mathbb)
        normal: 是否正体（覆盖 italic）

    返回:
        m:r XML 元素
    """
    r = _m_elem('r')
    rPr = _m_elem('rPr')

    if normal:
        nor = _m_elem('nor')
        rPr.append(nor)
    elif bold:
        sty = _m_elem('sty')
        sty.set(_m_qn('val'), 'b')
        rPr.append(sty)
    elif script:
        scr = _m_elem('scr')
        scr.set(_m_qn('val'), script)
        rPr.append(scr)
    elif not italic:
        nor = _m_elem('nor')
        rPr.append(nor)

    # 仅在 rPr 有子元素时追加
    if len(rPr) > 0:
        r.append(rPr)

    t_elem = _m_elem('t')
    t_elem.set(qn('xml:space'), 'preserve')
    t_elem.text = text
    r.append(t_elem)
    return r


def _build_matrix(cell_ommls, left_delim=None, right_delim=None):
    """从已解析的单元格 OMML 元素构建矩阵结构

    参数:
        cell_ommls: 二维列表 [[cell1, cell2, ...], [cell1, cell2, ...], ...]
                    每个 cell 是 OMML 元素列表（由 _LatexParser.parse() 返回）
        left_delim: 左定界符字符（可选，如 '(' '[' '|' '{'）
        right_delim: 右定界符字符（可选，如 ')' ']' '|' '}'）

    返回:
        m:d 或 m:m XML 元素
    """
    ncols = max((len(row) for row in cell_ommls), default=0)

    mat = _m_elem('m')
    mPr = _m_elem('mPr')
    mcs = _m_elem('mcs')
    mc = _m_elem('mc')
    mcPr = _m_elem('mcPr')
    count = _m_elem('count')
    count.set(_m_qn('val'), str(ncols))
    mcPr.append(count)
    mcJc = _m_elem('mcJc')
    mcJc.set(_m_qn('val'), 'center')
    mcPr.append(mcJc)
    mc.append(mcPr)
    mcs.append(mc)
    mPr.append(mcs)
    mat.append(mPr)

    # 构建行
    for row in cell_ommls:
        mr = _m_elem('mr')
        for cell_children in row:
            e = _m_elem('e')
            for child in cell_children:
                if isinstance(child, str):
                    e.append(_m_run(child))
                elif child is not None:
                    e.append(child)
            mr.append(e)
        mat.append(mr)

    # 有定界符时外层包裹 m:d
    if left_delim is not None or right_delim is not None:
        d = _m_elem('d')
        dPr = _m_elem('dPr')
        if left_delim is not None:
            begChr = _m_elem('begChr')
            begChr.set(_m_qn('val'), left_delim)
            dPr.append(begChr)
        if right_delim is not None:
            endChr = _m_elem('endChr')
            endChr.set(_m_qn('val'), right_delim)
            dPr.append(endChr)
        # 只有左定界符时也需显式声明右定界符为空，否则 Word 自动配对
        if left_delim is not None and right_delim is None:
            endChr = _m_elem('endChr')
            endChr.set(_m_qn('val'), '')
            dPr.append(endChr)
        d.append(dPr)
        e = _m_elem('e')
        e.append(mat)
        d.append(e)
        return d

    return mat


# 矩阵环境 → 定界符映射
_MATRIX_DELIMS = {
    'matrix': (None, None),
    'bmatrix': ('[', ']'),
    'pmatrix': ('(', ')'),
    'vmatrix': ('|', '|'),
    'Vmatrix': ('‖', '‖'),
}


# 字体样式命令 → _m_run_with_style 参数映射
_FONT_STYLES = {
    'mathbf':      {'bold': True},
    'bm':          {'bold': True},
    'boldsymbol':  {'bold': True},
    'mathit':      {'italic': True},
    'mathrm':      {'normal': True},
    'mathcal':     {'script': 'script'},
    'mathscr':     {'script': 'script'},
    'mathbb':      {'script': 'double-struck'},
    'mathtt':      {'italic': False},  # 等宽正体
}


def _apply_style_to_elem(elem, style_args):
    """递归对 OMML 元素树中的所有 m:r 元素应用字体样式

    参数:
        elem: OMML XML 元素
        style_args: 传给 _m_run_with_style 的关键字参数字典

    返回:
        应用样式后的元素（可能是新元素）
    """
    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
    if tag == 'r':
        # 替换 m:r 元素: 保留文本内容，用 _m_run_with_style 重新生成
        t_elem = elem.find(f'{{{NSM}}}t')
        text = t_elem.text if t_elem is not None else ''
        italic = not style_args.get('normal', False) and style_args.get('italic', True)
        bold = style_args.get('bold', False)
        script = style_args.get('script', None)
        normal = style_args.get('normal', False)
        new_r = _m_run_with_style(text, italic=italic, bold=bold, script=script, normal=normal)
        # 在父元素中替换
        parent = elem.getparent()
        if parent is not None:
            idx = list(parent).index(elem)
            parent[idx] = new_r
        return new_r
    else:
        for child in list(elem):
            _apply_style_to_elem(child, style_args)
        return elem


def _try_merge_text_runs(children):
    """若 children 全为简单 m:r 文本元素，返回合并后的文本；否则返回 None"""
    texts = []
    for child in children:
        if isinstance(child, str):
            texts.append(child)
            continue
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'r':
            return None
        t_elem = child.find(f'{{{NSM}}}t')
        if t_elem is None:
            return None
        texts.append(t_elem.text or '')
    return ''.join(texts)


class _LatexParser:
    """LaTeX 数学公式递归下降解析器，生成 OMML XML"""

    def __init__(self, latex, in_env=False, display=False):
        self.s = latex
        self.pos = 0
        self._in_env = in_env  # 是否在矩阵/环境内，影响 & 和 \\ 的处理
        self._display = display  # 是否行间公式，影响 \frac 大小

    def peek(self):
        if self.pos < len(self.s):
            return self.s[self.pos]
        return None

    def consume(self):
        c = self.peek()
        if c is not None:
            self.pos += 1
        return c

    def expect(self, expected):
        if self.peek() != expected:
            return False
        self.consume()
        return True

    def skip_spaces(self):
        while self.peek() in (' ', '\t', '\n'):
            self.consume()

    def read_name(self):
        """读取反斜杠后的命令名，如 \frac → 'frac'"""
        name = []
        while self.peek() and self.peek().isalpha():
            name.append(self.consume())
        return ''.join(name)

    def read_required_arg(self):
        """读取花括号中的必选参数 { ... }，注意嵌套"""
        self.skip_spaces()
        if not self.expect('{'):
            return None
        depth = 1
        buf = []
        while self.pos < len(self.s) and depth > 0:
            c = self.consume()
            if c == '{':
                depth += 1
            elif c == '}':
                depth -= 1
                if depth == 0:
                    break
            buf.append(c)
        return ''.join(buf)

    def _scan_until_end(self, env_name):
        """扫描原始文本直到匹配的 \\end{env_name}，处理嵌套 \\begin{env_name}。

        返回:
            原始文本内容（不含 \\begin 和 \\end 标记）
        """
        body_chars = []
        depth = 1
        while self.pos < len(self.s) and depth > 0:
            c = self.peek()
            if c == '\\':
                saved = self.pos
                self.consume()  # consume '\'
                name = self.read_name()
                if name == 'begin':
                    inner_env = self.read_required_arg()
                    if inner_env == env_name:
                        depth += 1
                    body_chars.append(f'\\begin{{{inner_env}}}')
                elif name == 'end':
                    inner_env = self.read_required_arg()
                    if inner_env == env_name:
                        depth -= 1
                        if depth == 0:
                            break
                    body_chars.append(f'\\end{{{inner_env}}}')
                else:
                    # 非 begin/end 命令，视为普通内容
                    body_chars.append('\\' + name)
                    if self.peek() == '{':
                        body_chars.append('{' + (self.read_required_arg() or '') + '}')
            else:
                body_chars.append(self.consume())
        return ''.join(body_chars)

    def _parse_environment(self, env_name):
        """解析 \\begin{env_name}...\\end{env_name} 并生成 OMML 元素

        支持的 environment:
            - matrix 系列: matrix, bmatrix, pmatrix, vmatrix, Vmatrix
            - cases: 分段函数
            - align / align* / aligned: 多行对齐

        返回:
            OMML XML 元素，或（align 环境）m:oMathPara 元素 / m:oMath 列表
        """
        body = self._scan_until_end(env_name)

        if env_name == 'cases':
            return self._parse_cases_env(body)

        if env_name in _MATRIX_DELIMS:
            return self._parse_matrix_env(body, env_name)

        # align 环境
        if env_name in ('align', 'align*', 'aligned'):
            return self._parse_align_lines(body)

        # 未知环境：回退为文本
        return _m_run(f'\\begin{{{env_name}}}{body}\\end{{{env_name}}}')

    def _parse_align_lines(self, body):
        """解析 align 环境体：每行一个公式，& 为对齐点（在 OMML 中移除）

        返回:
            m:oMathPara 元素（单行时返回 m:oMath）
        """
        # 按 \\ 分行
        raw_lines = body.split('\\\\')
        omml_lines = []
        for line in raw_lines:
            # 移除 & 对齐标记并清理多余空格
            # 例如: "x + y &= 2" → "x + y = 2"
            text = re.sub(r'\s*&\s*', '', line).strip()
            if text:
                children = _LatexParser(text, in_env=True).parse()
                omml = _m_elem('oMath')
                for c in children:
                    omml.append(c if not isinstance(c, str) else _m_run(c))
                omml_lines.append(omml)

        if not omml_lines:
            return _m_run('')
        if len(omml_lines) == 1:
            return omml_lines[0]
        # 多行：m:oMathPara 包裹
        omp = _m_elem('oMathPara')
        # 设置段落属性：居中对齐
        ompPr = _m_elem('oMathParaPr')
        jc = _m_elem('jc')
        jc.set(_m_qn('val'), 'center')
        ompPr.append(jc)
        omp.append(ompPr)
        for omml in omml_lines:
            omp.append(omml)
        return omp

    def _parse_matrix_env(self, body, env_name):
        """解析矩阵环境体并生成 OMML 矩阵元素

        body 按 \\\\ 分行，每行按 & 分列，每个单元格用 _LatexParser 解析。
        """
        # 按 \\ 分行
        row_texts = body.split('\\\\')
        cell_ommls = []
        for row_text in row_texts:
            # 每行按 & 分列
            col_texts = row_text.split('&')
            row_cells = []
            for col_text in col_texts:
                col_text = col_text.strip()
                if col_text:
                    cell_children = _LatexParser(col_text, in_env=True).parse()
                else:
                    cell_children = []
                row_cells.append(cell_children)
            if row_cells:
                cell_ommls.append(row_cells)

        if not cell_ommls:
            return _m_run('')

        left_delim, right_delim = _MATRIX_DELIMS[env_name]
        return _build_matrix(cell_ommls, left_delim, right_delim)

    def _parse_cases_env(self, body):
        """解析 cases 环境体：左花括号定界符 + 内部矩阵

        每行格式: 表达式 & 条件
        """
        row_texts = body.split('\\\\')
        cell_ommls = []
        for row_text in row_texts:
            col_texts = row_text.split('&')
            row_cells = []
            for col_text in col_texts:
                col_text = col_text.strip()
                if col_text:
                    cell_children = _LatexParser(col_text, in_env=True).parse()
                else:
                    cell_children = []
                row_cells.append(cell_children)
            # 确保每行至少 2 列（cases 格式）
            while len(row_cells) < 2:
                row_cells.append([])
            if row_cells:
                cell_ommls.append(row_cells)

        if not cell_ommls:
            return _m_run('')

        # cases: 左花括号，无右定界符（设空字符串禁用默认配对标点）
        return _build_matrix(cell_ommls, '{', '')

    def parse_expr(self):
        """解析表达式：term*。环境内遇到 & 或 \\\\ 时停止。"""
        children = []
        while self.pos < len(self.s):
            c = self.peek()
            if c is None or c == '}':
                break
            if c == '&':  # 环境内的列分隔，停止解析
                if self._in_env:
                    break
                self.consume()
                children.append(_m_run('&'))
                continue
            if c == '\\' and self.pos + 1 < len(self.s) and self.s[self.pos + 1] == '\\':
                # \\ 行分隔：环境内停止，环境外当文本
                if self._in_env:
                    break
                self.consume()
                self.consume()
                children.append(_m_run('\\\\'))
                continue
            term = self.parse_term()
            if term is not None:
                children.append(term)
            if self.peek() == '}':
                break
        return children

    def parse_term(self):
        """解析单个 term：atom 后可选 sub/sup"""
        c = self.peek()
        if c is None or c == '}':
            return None

        # 上标 ^
        if c == '^':
            self.consume()
            sup_elem = self.parse_atom()
            if sup_elem is not None:
                sSup = _m_elem('sSup')
                sSup_e = _m_elem('e')
                sSup_e.append(_m_run(''))  # 空基数
                sSup.append(sSup_e)
                sSup_sup = _m_elem('sup')
                sSup_sup.append(sup_elem if not isinstance(sup_elem, str) else _m_run(sup_elem))
                sSup.append(sSup_sup)
                return sSup  # Don't try to attach sub/sup further
            else:
                return _m_run('^')

        # 下标 _
        if c == '_':
            self.consume()
            sub_elem = self.parse_atom()
            if sub_elem is not None:
                sSub = _m_elem('sSub')
                sSub_e = _m_elem('e')
                sSub_e.append(_m_run(''))
                sSub.append(sSub_e)
                sSub_sub = _m_elem('sub')
                sSub_sub.append(sub_elem if not isinstance(sub_elem, str) else _m_run(sub_elem))
                sSub.append(sSub_sub)
                return sSub
            else:
                return _m_run('_')

        # atom
        atom = self.parse_atom()
        if atom is None:
            return None

        # 检查后续的 sub/sup
        c = self.peek()

        # _ 后接 ^ (如 x_i^2 或 \sum_{i=1}^{n})
        if c == '_':
            self.consume()
            sub_arg = self.parse_atom()
            sup_arg = None
            if self.peek() == '^':
                self.consume()
                sup_arg = self.parse_atom()

            if sup_arg is not None:
                sSubSup = _m_elem('sSubSup')
                e = _m_elem('e')
                e.append(atom if not isinstance(atom, str) else _m_run(atom))
                sSubSup.append(e)
                s = _m_elem('sub')
                s.append(sub_arg if not isinstance(sub_arg, str) else _m_run(sub_arg))
                sSubSup.append(s)
                sup = _m_elem('sup')
                sup.append(sup_arg if not isinstance(sup_arg, str) else _m_run(sup_arg))
                sSubSup.append(sup)
                return sSubSup
            else:
                sSub = _m_elem('sSub')
                e = _m_elem('e')
                e.append(atom if not isinstance(atom, str) else _m_run(atom))
                sSub.append(e)
                s = _m_elem('sub')
                s.append(sub_arg if not isinstance(sub_arg, str) else _m_run(sub_arg))
                sSub.append(s)
                return sSub

        elif c == '^':
            self.consume()
            sup_arg = self.parse_atom()
            if sup_arg is not None:
                sSup = _m_elem('sSup')
                e = _m_elem('e')
                e.append(atom if not isinstance(atom, str) else _m_run(atom))
                sSup.append(e)
                sup = _m_elem('sup')
                sup.append(sup_arg if not isinstance(sup_arg, str) else _m_run(sup_arg))
                sSup.append(sup)
                return sSup

        return atom

    def parse_atom(self):
        """解析最小单元：字母/数字、命令、花括号组"""
        c = self.peek()
        if c is None or c == '}' or c == '&':
            return None

        # 花括号组 { ... }
        if c == '{':
            self.consume()
            children = self.parse_expr()
            self.expect('}')  # consume closing brace if present
            if len(children) == 1:
                return children[0]
            # 多子元素时检查是否全为普通文本 run，若是则合并避免冗余 m:d 包装
            merged_text = _try_merge_text_runs(children)
            if merged_text is not None:
                return _m_run(merged_text)
            return _m_wrap_d(children)

        # LaTeX 命令 \
        if c == '\\':
            self.consume()
            if self.peek() is None:
                return _m_run('\\')

            # 检查是否为空
            if self.peek() in (' ', '\t', '\n'):
                return _m_run(' ')

            # LaTeX 间距字符: \, \! \: \;
            if self.peek() == ',':
                self.consume()
                return _m_run(' ')  # thin space U+2009
            if self.peek() == '!':
                self.consume()
                return _m_run('')  # negative thin space
            if self.peek() == ':':
                self.consume()
                return _m_run(' ')  # medium space (en space)

            name = self.read_name()
            if not name:
                # 转义字符如 \$ \# \% \& \_ \{ \}
                esc = self.consume()
                if esc:
                    return _m_run(esc)
                return _m_run('\\')

            # 环境: \begin{env} ... \end{env}
            if name == 'begin':
                env_name = self.read_required_arg()
                if env_name:
                    return self._parse_environment(env_name)
                return _m_run('\\begin{}')
            if name == 'end':
                # \end{env} 由 _parse_environment 处理，不应直接出现在 parse_atom 中
                # 如果出现（解析异常），回退为文本
                env_name = self.read_required_arg()
                return _m_run(f'\\end{{{env_name or ""}}}')

            # 字体样式 \mathbf{}, \mathbb{}, \mathcal{} 等
            if name in _FONT_STYLES:
                arg_text = self.read_required_arg()
                if arg_text is not None:
                    style_args = _FONT_STYLES[name]
                    # 解析参数内容
                    inner_children = _LatexParser(arg_text).parse()
                    # 对每个结果元素应用样式
                    styled = []
                    for child in inner_children:
                        if isinstance(child, str):
                            styled.append(child)
                        else:
                            styled.append(_apply_style_to_elem(child, style_args))
                    if len(styled) == 1 and not isinstance(styled[0], str):
                        return styled[0]
                    # 尝试合并纯文本 run，避免多余 m:d 包装
                    merged = _try_merge_text_runs(styled)
                    if merged is not None:
                        # 用合并后的文本 + 相同样式重新生成
                        italic = not style_args.get('normal', False) and style_args.get('italic', True)
                        return _m_run_with_style(merged, italic=italic,
                                                 bold=style_args.get('bold', False),
                                                 script=style_args.get('script'),
                                                 normal=style_args.get('normal', False))
                    return _m_wrap_d(styled)
                return _m_run('\\' + name + '{}')

            # 间距命令 \, \! \quad \qquad \enspace
            if name == ',' or name == 'thinspace':
                return _m_run(' ')  # thin space U+2009
            if name == '!' or name == 'negthinspace':
                return _m_run('')  # 零宽负间距
            if name == 'quad':
                return _m_run(' ')  # em space U+2003
            if name == 'qquad':
                return _m_run('  ')  # 2em
            if name == 'enspace':
                return _m_run(' ')  # en space U+2002

            # 扩展重音 \overrightarrow{...} \overleftarrow{...}
            if name == 'overrightarrow':
                body = self.read_required_arg()
                if body is not None:
                    acc = _m_elem('acc')
                    accPr = _m_elem('accPr')
                    chr_elem = _m_elem('chr')
                    chr_elem.set(_m_qn('val'), '→')
                    accPr.append(chr_elem)
                    acc.append(accPr)
                    e = _m_elem('e')
                    e.append(_m_run(body))
                    acc.append(e)
                    return acc
                return _m_run('\\overrightarrow{}')
            if name == 'overleftarrow':
                body = self.read_required_arg()
                if body is not None:
                    acc = _m_elem('acc')
                    accPr = _m_elem('accPr')
                    chr_elem = _m_elem('chr')
                    chr_elem.set(_m_qn('val'), '←')
                    accPr.append(chr_elem)
                    acc.append(accPr)
                    e = _m_elem('e')
                    e.append(_m_run(body))
                    acc.append(e)
                    return acc
                return _m_run('\\overleftarrow{}')

            # 颜色命令（简化处理：转为普通文本，颜色属性在 OMML 中支持有限）
            # \textcolor{color}{text} → 忽略颜色，保留文本
            if name == 'textcolor':
                color_name = self.read_required_arg()
                text_arg = self.read_required_arg()
                if text_arg is not None:
                    return _m_text_run(text_arg)
                return _m_run('\\textcolor{}{}')
            if name == 'color':
                # \color{name} 设置后续颜色，无法精确映射到 OMML，忽略
                _ = self.read_required_arg()
                return _m_run('')  # 忽略颜色标记

            # n-ary 大运算符（sum, prod, int 等）
            if name in BIG_OPS:
                # 先检查是否有下上限
                saved = self.pos
                has_limits = False
                self.skip_spaces()
                if self.peek() == '_' or self.peek() == '^':
                    has_limits = True
                self.pos = saved  # 回退

                # 无上下限：直接输出 Unicode 字符，避免 m:nary 占位符
                if not has_limits:
                    return _m_run(BIG_OPS[name])

                nary = _m_elem('nary')
                naryPr = _m_elem('naryPr')
                chr_elem = _m_elem('chr')
                chr_elem.set(_m_qn('val'), BIG_OPS[name])
                naryPr.append(chr_elem)
                limLoc = _m_elem('limLoc')
                limLoc.set(_m_qn('val'), 'undOvr')  # limits above and below
                naryPr.append(limLoc)
                nary.append(naryPr)

                # 记录是否有上下限，无则隐藏占位符
                has_sub = False
                has_sup = False

                # 读取下上限
                if self.peek() == '_':
                    self.consume()
                    has_sub = True
                    sub_arg = self.parse_atom()
                    if sub_arg is not None:
                        sub = _m_elem('sub')
                        sub.append(sub_arg if not isinstance(sub_arg, str) else _m_run(sub_arg))
                        nary.append(sub)

                if self.peek() == '^':
                    self.consume()
                    has_sup = True
                    sup_arg = self.parse_atom()
                    if sup_arg is not None:
                        sup = _m_elem('sup')
                        sup.append(sup_arg if not isinstance(sup_arg, str) else _m_run(sup_arg))
                        nary.append(sup)

                # 无上下限时隐藏占位符
                if not has_sub:
                    subHide = _m_elem('subHide')
                    subHide.set(_m_qn('val'), '1')
                    naryPr.append(subHide)
                if not has_sup:
                    supHide = _m_elem('supHide')
                    supHide.set(_m_qn('val'), '1')
                    naryPr.append(supHide)

                e = _m_elem('e')
                # 读取 integrand / summand
                rest = self.parse_expr()
                if rest:
                    for r in rest:
                        e.append(r if not isinstance(r, str) else _m_run(r))
                nary.append(e)
                return nary

            # 分数 \frac{num}{den} 及展示变体 \dfrac \tfrac
            if name in ('frac', 'dfrac', 'tfrac'):
                num_text = self.read_required_arg()
                den_text = self.read_required_arg()
                if num_text is not None and den_text is not None:
                    f = _m_elem('f')
                    # fPr: 分数类型区分显示/行内/小分数
                    fPr = _m_elem('fPr')
                    type_elem = _m_elem('type')
                    if name == 'dfrac' or (name == 'frac' and self._display):
                        type_elem.set(_m_qn('val'), 'bar')   # 显示分数（大）
                    elif name == 'tfrac' or name == 'frac':
                        type_elem.set(_m_qn('val'), 'skw')   # 小分数（行内）
                    fPr.append(type_elem)
                    f.append(fPr)
                    num = _m_elem('num')
                    num_parsed = _LatexParser(num_text).parse_expr()
                    for item in num_parsed:
                        num.append(item if not isinstance(item, str) else _m_run(item))
                    f.append(num)
                    den = _m_elem('den')
                    den_parsed = _LatexParser(den_text).parse_expr()
                    for item in den_parsed:
                        den.append(item if not isinstance(item, str) else _m_run(item))
                    f.append(den)
                    return f
                return _m_run(f'\\frac{{{num_text}}}{{{den_text}}}')

            # 平方根 \sqrt{...} 或 \sqrt[n]{...}
            if name == 'sqrt':
                rad = _m_elem('rad')
                deg = None
                if self.peek() == '[':
                    self.consume()
                    deg_text = []
                    while self.peek() and self.peek() != ']':
                        deg_text.append(self.consume())
                    self.expect(']')
                    deg_text = ''.join(deg_text)
                    deg = _m_elem('deg')
                    deg.append(_m_run(deg_text))
                if deg is not None:
                    rad.append(deg)
                else:
                    # 无 [n] 时隐藏 degree 占位符（Word 默认显示空位）
                    radPr = _m_elem('radPr')
                    degHide = _m_elem('degHide')
                    degHide.set(_m_qn('val'), '1')
                    radPr.append(degHide)
                    rad.append(radPr)
                e = _m_elem('e')
                body = self.read_required_arg()
                if body is not None:
                    body_parsed = _LatexParser(body).parse_expr()
                    for item in body_parsed:
                        e.append(item if not isinstance(item, str) else _m_run(item))
                rad.append(e)
                return rad

            # 函数名 \sin, \cos, \log 等
            if name in FUNCTIONS:
                return _m_run(name)

            # 文本 \text{...}
            if name == 'text':
                txt = self.read_required_arg()
                return _m_text_run(txt) if txt else _m_run('\\text{}')

            # 重音符 \hat, \bar, \vec 等
            if name in ACCENTS:
                acc = _m_elem('acc')
                accPr = _m_elem('accPr')
                chr_elem = _m_elem('chr')
                chr_elem.set(_m_qn('val'), ACCENTS[name])
                accPr.append(chr_elem)
                acc.append(accPr)
                e = _m_elem('e')
                # 重音符可以跟花括号参数或单个字符
                base = self.parse_atom()
                if base is not None:
                    e.append(base if not isinstance(base, str) else _m_run(base))
                acc.append(e)
                return acc

            # 定界符 \left( \right) \left[ \right] 等
            if name == 'left':
                delim_char = self.consume()
                # 处理 \left. （不可见定界符）
                if delim_char == '.':
                    delim_char = ''
                d = _m_elem('d')
                dPr = _m_elem('dPr')
                begChr = _m_elem('begChr')
                begChr.set(_m_qn('val'), delim_char or '')
                dPr.append(begChr)
                d.append(dPr)
                # 读取直到 \right
                body = []
                while self.pos < len(self.s):
                    if self.peek() == '\\':
                        saved = self.pos
                        self.consume()
                        name2 = self.read_name()
                        if name2 == 'right':
                            right_delim = self.consume()
                            if right_delim == '.':
                                right_delim = ''
                            endChr = _m_elem('endChr')
                            endChr.set(_m_qn('val'), right_delim or '')
                            dPr.append(endChr)
                            break
                        else:
                            # 不是 \right，回退
                            self.pos = saved
                            body.append(self.parse_term())
                    else:
                        body.append(self.parse_term())
                # 内容必须包在 m:e 中
                e = _m_elem('e')
                for b in body:
                    if b is not None:
                        e.append(b if not isinstance(b, str) else _m_run(b))
                d.append(e)
                return d

            # 花括号（literal braces via \{ \}）
            if name == '{':
                return _m_run('{')
            if name == '}':
                return _m_run('}')

            # 希腊字母
            if name in GREEK:
                return _m_run(GREEK[name])

            # 符号
            if name in SYMBOLS:
                return _m_run(SYMBOLS[name])

            # 未知命令：保留原始文本
            result = ['\\' + name]
            if self.peek() == '{':
                result.append('{' + (self.read_required_arg() or '') + '}')
            return _m_run(''.join(result))

        # 普通字符
        ch = self.consume()
        return _m_run(ch)

    def parse(self):
        """主入口：解析整个 LaTeX 字符串并返回 OMML 元素列表"""
        self.skip_spaces()
        return self.parse_expr()


def latex_to_omml(latex, display=False):
    """将 LaTeX 数学公式转换为 OMML XML 元素

    参数:
        latex: LaTeX 公式字符串（不含 $ 定界符）
        display: True=行间公式(m:oMathPara), False=行内公式(m:oMath)
    返回:
        OMML 顶级元素（m:oMath 或 m:oMathPara），可直接插入段落 XML
    """
    parser = _LatexParser(latex, display=display)
    children = parser.parse()

    # 如果解析结果已经有 m:oMathPara（align 环境），直接返回
    if len(children) == 1:
        child = children[0]
        if hasattr(child, 'tag') and child.tag == f'{{{NSM}}}oMathPara':
            return child

    omml = _m_elem('oMath')
    for c in children:
        omml.append(c if not isinstance(c, str) else _m_run(c))

    if display:
        omp = _m_elem('oMathPara')
        omp.append(omml)
        return omp

    return omml


# ============================================================
# 字体工具
# ============================================================

def set_run_font(run, cn_font, en_font="Times New Roman", size_pt=10.5, bold=False):
    """设置 run 的中英文字体、字号、加粗"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = en_font
    run.font.color.rgb = None  # 黑色
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)


def _set_para_mark_size(paragraph, size_pt):
    """设置段落标记字号（半磅），供行内公式 OMML 继承

    参数：
        paragraph: python-docx Paragraph 对象
        size_pt: 字号，单位 pt（如 9 → 9pt）
    """
    pPr = paragraph._element.get_or_add_pPr()
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), str(int(size_pt * 2)))


def add_empty_para(doc):
    """添加五号空行"""
    p = doc.add_paragraph()
    run = p.add_run('')
    set_run_font(run, '宋体', size_pt=10.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


# ============================================================
# 图片工具
# ============================================================

def download_image(url):
    """下载或复制图片到临时文件，返回路径。失败返回 None。

    支持：
    - HTTP/HTTPS URL：requests 下载
    - 本地文件路径：直接复制到临时文件
    - base64 data URI：解码保存
    """
    import requests
    import shutil

    # 1. 本地文件路径（相对于工作目录的路径）
    if os.path.exists(url):
        suffix = os.path.splitext(url)[1] or '.png'
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        shutil.copy2(url, tmp.name)
        tmp.close()
        return tmp.name

    # 2. base64 data URI
    if url.startswith('data:'):
        import base64
        # 格式: data:image/png;base64,xxxx 或 data:image/png,xxxx
        header, data = url.split(',', 1)
        is_base64 = ';base64' in header
        suffix = '.png'
        mime_match = re.match(r'data:image/(\w+)', header)
        if mime_match:
            suffix = '.' + mime_match.group(1)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        tmp.write(base64.b64decode(data) if is_base64 else data.encode())
        tmp.close()
        return tmp.name

    # 3. HTTP/HTTPS URL
    headers = {'User-Agent': 'Mozilla/5.0 (compatible; mddoc-converter/1.0)'}
    try:
        resp = requests.get(url, headers=headers, timeout=30)
        resp.raise_for_status()
        suffix = os.path.splitext(urllib.parse.urlparse(url).path)[1] or '.png'
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        tmp.write(resp.content)
        tmp.close()
        return tmp.name
    except Exception:
        return None


def calc_image_size(img_path):
    """计算图片在 docx 中的显示尺寸：8-12cm 规则，等比缩放"""
    img = Image.open(img_path)
    dpi = img.info.get('dpi', (96, 96))
    dpi_x = dpi[0] if dpi and dpi[0] > 0 else 96
    dpi_y = dpi[1] if dpi and dpi[1] > 0 else 96

    width_cm = img.width / dpi_x * 2.54
    height_cm = img.height / dpi_y * 2.54

    if width_cm > 12:
        target_w = 12.0
    elif width_cm < 8:
        target_w = 8.0
    else:
        target_w = width_cm

    ratio = target_w / width_cm
    return Cm(target_w), Cm(height_cm * ratio)


def add_placeholder_image(doc, text):
    """生成占位图片：灰色矩形 + 文字描述"""
    from PIL import ImageDraw, ImageFont
    img = Image.new('RGB', (600, 200), (220, 220, 220))
    draw = ImageDraw.Draw(img)
    draw.text((10, 80), text, fill=(80, 80, 80))
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    img.save(tmp, format='PNG')
    tmp.close()
    return tmp.name


# ============================================================
# 表格工具
# ============================================================

def set_three_line_table(table):
    """三线表：仅顶线粗+底线粗，表头格底线细，数据行之间无线，无竖线"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    borders = OxmlElement('w:tblBorders')

    # 顶线：粗 1.5pt
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '000000')
    borders.append(top)

    # 底线：粗 1.5pt
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    borders.append(bottom)

    # 所有内部线+竖线全部关闭（不用insideH！数据行之间无线）
    for edge in ('left', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'none')
        elem.set(qn('w:sz'), '0')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), 'auto')
        borders.append(elem)

    tblPr.append(borders)

    # 表头行每格底部单独加细线 0.75pt（cell级border，非table级insideH）
    if table.rows:
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            cell_bottom = OxmlElement('w:bottom')
            cell_bottom.set(qn('w:val'), 'single')
            cell_bottom.set(qn('w:sz'), '6')
            cell_bottom.set(qn('w:space'), '0')
            cell_bottom.set(qn('w:color'), '000000')
            tcBorders.append(cell_bottom)
            tcPr.append(tcBorders)


def add_page_break_before(paragraph):
    """在段落前插入分页符"""
    pPr = paragraph._element.get_or_add_pPr()
    pb = OxmlElement('w:pageBreakBefore')
    pPr.append(pb)


def _set_paragraph_left_chars(paragraph, chars=2):
    """设置段落左缩进（字符单位），通过 XML w:ind 的 w:leftChars 属性

    参数：
        paragraph: python-docx Paragraph 对象
        chars: 缩进字符数（默认 2）
    """
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.insert(0, ind)
    ind.set(qn('w:leftChars'), str(chars * 100))


def set_first_line_indent_chars(paragraph, chars=2):
    """设置段落首行缩进（字符单位），通过 XML w:ind 的 firstLineChars 属性

    参数：
        paragraph: python-docx Paragraph 对象
        chars: 缩进字符数（0 表示无缩进，默认 2）
    """
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        # 插入到 pPr 的第一个子元素之前（通常在 rPr 之前）
        pPr.insert(0, ind)
    if chars > 0:
        ind.set(qn('w:firstLineChars'), str(chars * 100))
    else:
        # 清除首行缩进
        for attr in ('w:firstLineChars', 'w:firstLine'):
            key = qn(attr)
            if key in ind.attrib:
                del ind.attrib[key]


def set_outline_level(paragraph, level):
    """设置 outline level（XML方式，兼容所有python-docx版本）"""
    pPr = paragraph._element.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl')
    ol.set(qn('w:val'), str(level))
    pPr.append(ol)


def override_builtin_heading_styles(doc):
    """覆盖 Word 内置 Heading 1-6 样式为学术格式，防止默认蓝色/Cambria/加粗干扰

    参数：
        doc: python-docx Document 对象
    """
    specs = {
        'Heading 1': ('黑体', 16, True, WD_ALIGN_PARAGRAPH.CENTER),
        'Heading 2': ('黑体', 14, False, WD_ALIGN_PARAGRAPH.LEFT),
        'Heading 3': ('宋体', 12, False, WD_ALIGN_PARAGRAPH.LEFT),
        'Heading 4': ('宋体', 10.5, False, WD_ALIGN_PARAGRAPH.LEFT),
        'Heading 5': ('宋体', 10.5, False, WD_ALIGN_PARAGRAPH.LEFT),
        'Heading 6': ('宋体', 10.5, False, WD_ALIGN_PARAGRAPH.LEFT),
    }
    for name, (cn_font, size_pt, bold, align) in specs.items():
        style = doc.styles[name]
        style.font.size = Pt(size_pt)
        style.font.name = 'Times New Roman'
        style.font.bold = bold
        style.font.color.rgb = None
        style.font.italic = False
        style.font.underline = False
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), cn_font)
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        pPr = style.element.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '312')
        spacing.set(qn('w:lineRule'), 'auto')

        # 缩进全部显式置零（twips 与 chars 双维度），
        # 防止 Word/WPS 对 outlineLvl 段落套用内置标题样式自带的首行缩进
        ind = OxmlElement('w:ind')
        for attr in ('w:left', 'w:leftChars', 'w:firstLine', 'w:firstLineChars',
                     'w:hanging', 'w:hangingChars'):
            ind.set(qn(attr), '0')
        ol = pPr.find(qn('w:outlineLvl'))
        if ol is not None:
            ol.addprevious(ind)
        else:
            pPr.append(ind)


# ============================================================
# Markdown 解析
# ============================================================

def parse_markdown(text):
    """使用 mistune 解析 Markdown 文本为 AST 节点列表

    参数：
        text: Markdown 原始文本

    返回:
        mistune AST 节点列表，包含 heading/paragraph/list/table/block_math/
        block_code/thematic_break 等类型
    """
    md = mistune.create_markdown(
        renderer='ast',
        plugins=[mistune_table.table, mistune_math.math]
    )
    return md(text)


# ============================================================
# DOCX 生成
# ============================================================

def _extract_ast_text(children):
    """从 mistune inline AST children 递归提取纯文本

    参数：
        children: mistune inline AST 子节点列表

    返回:
        合并后的纯文本字符串
    """
    parts = []
    for c in children:
        if c['type'] == 'text':
            parts.append(c.get('raw', ''))
        elif 'children' in c:
            parts.append(_extract_ast_text(c['children']))
    return ''.join(parts)


def _apply_omml_style(omml, bold=False, italic=False):
    """为 OMML 内所有 m:r 应用加粗/斜体样式（m:sty）

    参数：
        omml: m:oMath 元素
        bold: 是否加粗
        italic: 是否斜体（默认数学即斜体，仅强加时设置）

    返回:
        应用样式后的 omml（原地修改）
    """
    if not (bold or italic):
        return omml
    val = ('b' if bold else '') + ('i' if italic else '')
    for r in omml.iter(f'{{{NSM}}}r'):
        rPr = r.find(f'{{{NSM}}}rPr')
        if rPr is None:
            rPr = _m_elem('rPr')
            r.insert(0, rPr)
        if rPr.find(f'{{{NSM}}}sty') is None:
            sty = _m_elem('sty')
            sty.set(_m_qn('val'), val)
            rPr.append(sty)
    return omml


def walk_inline(paragraph, children, base_cn='宋体', base_en='Times New Roman',
                base_size=10.5, bold=False, italic=False, underline=False):
    """递归遍历 inline AST，将格式化文本和 OMML 公式附加到段落

    参数：
        paragraph: python-docx Paragraph 对象
        children: mistune inline AST 子节点列表
        base_cn: 中文字体名（默认 '宋体'）
        base_en: 英文字体名（默认 'Times New Roman'）
        base_size: 基础字号，单位 pt（默认 10.5）
        bold: 加粗状态（嵌套 strong 递归传 True）
        italic: 斜体状态（嵌套 emphasis 递归传 True）
        underline: 下划线状态（嵌套 link 递归传 True）
    """
    for child in children:
        ct = child['type']

        if ct == 'text':
            raw = child.get('raw', '')
            if raw:
                run = paragraph.add_run(raw)
                set_run_font(run, base_cn, en_font=base_en, size_pt=base_size, bold=bold)
                if italic:
                    run.font.italic = True
                if underline:
                    run.font.underline = True

        elif ct == 'strong':
            walk_inline(paragraph, child.get('children', []), base_cn, base_en, base_size,
                        bold=True, italic=italic, underline=underline)

        elif ct == 'emphasis':
            walk_inline(paragraph, child.get('children', []), base_cn, base_en, base_size,
                        bold=bold, italic=True, underline=underline)

        elif ct == 'inline_math':
            try:
                omml = latex_to_omml(child['raw'], display=False)
                omml = _apply_omml_style(omml, bold, italic)
                paragraph._element.append(omml)
            except Exception:
                run = paragraph.add_run(child.get('raw', ''))
                set_run_font(run, base_cn, en_font=base_en, size_pt=base_size, bold=bold)
                run.font.italic = True

        elif ct == 'codespan':
            run = paragraph.add_run(child.get('raw', ''))
            set_run_font(run, base_cn, en_font=base_en, size_pt=base_size)

        elif ct == 'link':
            walk_inline(paragraph, child.get('children', []), base_cn, base_en, base_size,
                        bold=bold, italic=italic, underline=True)

        elif ct == 'image':
            url = child.get('attrs', {}).get('url', '')
            alt = child.get('alt', '')
            img_path = download_image(url)
            if img_path is None:
                img_path = add_placeholder_image(None, alt or url)
            if img_path:
                w, h = calc_image_size(img_path)
                run = paragraph.add_run()
                run.add_picture(img_path, width=w)
                try:
                    os.unlink(img_path)
                except Exception:
                    pass

        elif ct == 'linebreak':
            run = paragraph.add_run()
            run.add_break()


def _ensure_list_numbering_defs(doc):
    """确保文档 numbering.xml 中存在有序和无序列表的抽象编号定义。
    python-docx 默认不含列表定义，需手动在 XML 级别注入。

    参数：
        doc: python-docx Document 对象
    返回：
        set: 当前已占用的 numId 集合（含有序 50 / 无序 51）
    """
    # 使用不与 Word 默认模板冲突的 numId 和 abstractNumId
    ORDERED_NUM_ID = '50'
    ORDERED_ABSTRACT_ID = '50'
    UNORDERED_NUM_ID = '51'
    UNORDERED_ABSTRACT_ID = '51'

    numbering_part = doc.part.numbering_part
    numbering_elem = numbering_part._element

    # 收集当前已占用的 numId
    used_ids = {num.get(qn('w:numId')) for num in numbering_elem.findall(qn('w:num'))}

    # 列表缩进由段落 leftChars 控制，numbering 不再叠加缩进
    BASE_INDENT = 0  # twips

    # 抽象编号 50：有序列表 (1, 2, 3..., a, b, c..., i, ii, iii...)
    abs_ord = OxmlElement('w:abstractNum')
    abs_ord.set(qn('w:abstractNumId'), ORDERED_ABSTRACT_ID)
    ord_fmts = [('decimal', '%1.'), ('lowerLetter', '%2)'), ('lowerRoman', '%3.')]
    for ilvl_val, (fmt, lvl_text) in enumerate(ord_fmts):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(ilvl_val))
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), fmt)
        lvl.append(numFmt)
        lvlText = OxmlElement('w:lvlText')
        lvlText.set(qn('w:val'), lvl_text)
        lvl.append(lvlText)
        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        left_val = BASE_INDENT * (ilvl_val + 1)
        ind.set(qn('w:left'), str(left_val))
        ind.set(qn('w:hanging'), str(BASE_INDENT))
        pPr.append(ind)
        lvl.append(pPr)
        abs_ord.append(lvl)
    numbering_elem.append(abs_ord)

    # 抽象编号 51：无序列表 (•, ◦, ▪)
    abs_unord = OxmlElement('w:abstractNum')
    abs_unord.set(qn('w:abstractNumId'), UNORDERED_ABSTRACT_ID)
    bull_chars = ['•', '◦', '▪']
    for ilvl_val, bull_char in enumerate(bull_chars):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(ilvl_val))
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), 'bullet')
        lvl.append(numFmt)
        lvlText = OxmlElement('w:lvlText')
        lvlText.set(qn('w:val'), bull_char)
        lvl.append(lvlText)
        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        left_val = BASE_INDENT * (ilvl_val + 1)
        ind.set(qn('w:left'), str(left_val))
        ind.set(qn('w:hanging'), str(BASE_INDENT))
        pPr.append(ind)
        lvl.append(pPr)
        abs_unord.append(lvl)
    numbering_elem.append(abs_unord)

    # num 实例：有序列表 → numId=50（仅当不存在时注入）
    if ORDERED_NUM_ID not in used_ids:
        num_ord = OxmlElement('w:num')
        num_ord.set(qn('w:numId'), ORDERED_NUM_ID)
        absRef_ord = OxmlElement('w:abstractNumId')
        absRef_ord.set(qn('w:val'), ORDERED_ABSTRACT_ID)
        num_ord.append(absRef_ord)
        numbering_elem.append(num_ord)
        used_ids.add(ORDERED_NUM_ID)

    # num 实例：无序列表 → numId=51（仅当不存在时注入）
    if UNORDERED_NUM_ID not in used_ids:
        num_unord = OxmlElement('w:num')
        num_unord.set(qn('w:numId'), UNORDERED_NUM_ID)
        absRef_unord = OxmlElement('w:abstractNumId')
        absRef_unord.set(qn('w:val'), UNORDERED_ABSTRACT_ID)
        num_unord.append(absRef_unord)
        numbering_elem.append(num_unord)
        used_ids.add(UNORDERED_NUM_ID)

    return used_ids


def _new_list_num_id(numbering_elem, abstract_id, used_ids):
    """为列表块创建独立的 w:num 实例并返回新的 numId。

    有序列表被其他内容（正文段落等）打断后需重新从 1 编号，
    因此每个有序列表块使用独立 numId（同一 abstractNum 的不同
    实例各自独立计数），而不再所有列表共享 numId=50。

    参数：
        numbering_elem: numbering.xml 根元素
        abstract_id: 引用的 abstractNumId（如 '50'）
        used_ids: 已占用 numId 集合，新分配的 numId 会加入其中
    返回：
        str: 新分配的 numId
    """
    nid = 52
    while str(nid) in used_ids:
        nid += 1
    nid_str = str(nid)
    used_ids.add(nid_str)
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), nid_str)
    abs_ref = OxmlElement('w:abstractNumId')
    abs_ref.set(qn('w:val'), abstract_id)
    num.append(abs_ref)
    numbering_elem.append(num)
    return nid_str


def _set_list_num_pr(paragraph, num_id, level=0):
    """设置段落的 Word 原生列表编号属性 (w:numPr)

    参数：
        paragraph: python-docx Paragraph 对象
        num_id: 引用的 numId（有序列表为每块独立实例，无序固定 '51'）
        level: 列表嵌套层级（0=顶层）
    """
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), num_id)
    numPr.append(numId)
    pPr.append(numPr)


def _render_list(doc, node, depth=0, used_ids=None):
    """渲染 Word 原生列表（有序/无序），支持多级嵌套和富文本 inline

    每个有序列表块分配独立 numId，被正文打断后重新从 1 编号；
    无序列表共享 numId='51'（项目符号无需重置）。

    参数：
        doc: python-docx Document 对象
        node: mistune list AST 节点
        depth: 列表嵌套层级（0=顶层）
        used_ids: 已占用 numId 集合（来自 _ensure_list_numbering_defs）
    """
    ordered = node.get('attrs', {}).get('ordered', False) if 'attrs' in node else node.get('ordered', False)

    if ordered:
        # 每个有序列表块独立 numId，保证打断后重新编号
        num_id = _new_list_num_id(doc.part.numbering_part._element, '50', used_ids)
    else:
        num_id = '51'

    for item in node.get('children', []):
        if item['type'] != 'list_item':
            continue

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        # 列表段前缩进：每级 2 字符（通过 leftChars），段落本身无首行缩进
        _set_paragraph_left_chars(p, 2 * (depth + 1))

        _set_list_num_pr(p, num_id, level=depth)

        for child in item.get('children', []):
            if child['type'] == 'block_text':
                walk_inline(p, child.get('children', []))
            elif child['type'] == 'paragraph':
                walk_inline(p, child.get('children', []))
            elif child['type'] == 'list':
                # 嵌套子列表：递归渲染
                _render_list(doc, child, depth=depth + 1, used_ids=used_ids)


def _render_table(doc, node, tab_counter, chapter_path, has_chapter, caption=''):
    """渲染三线表，适配 mistune table AST。

    参数：
        doc: Document 对象
        node: mistune table AST 节点
        tab_counter: 表格计数 dict
        chapter_path: 章节路径 list
        has_chapter: 是否有章标题
        caption: 表题文本（来自"表 xxx"段落缓冲，或为空回退到 header_cells[0]）
    """
    children = node.get('children', [])
    head_cells = []
    body_rows = []

    for c in children:
        if c['type'] == 'table_head':
            # table_head 的子节点直接是 table_cell（无 row 包装）
            for cell in c.get('children', []):
                if cell['type'] == 'table_cell':
                    head_cells.append(cell.get('children', []))
        elif c['type'] == 'table_body':
            # table_body > table_row > table_cell
            for row in c.get('children', []):
                if row['type'] != 'table_row':
                    continue
                row_data = []
                for cell in row.get('children', []):
                    if cell['type'] == 'table_cell':
                        row_data.append(cell.get('children', []))
                body_rows.append(row_data)

    if not head_cells:
        return

    ncols = len(head_cells)

    # 表题文本：优先使用传入的 caption，否则回退到第一列表头
    if not caption:
        caption = _extract_ast_text(head_cells[0]) if head_cells else ''
    tab_key = tuple(chapter_path[:1]) if has_chapter else None
    if tab_key:
        tab_counter[tab_key] = tab_counter.get(tab_key, 0) + 1
        tab_num = tab_counter[tab_key]
    else:
        tab_counter[None] = tab_counter.get(None, 0) + 1
        tab_num = tab_counter[None]

    tab_label = f"表{chapter_path[0]}-{tab_num}" if has_chapter else f"表{tab_num}"

    add_empty_para(doc)
    p_tab_cap = doc.add_paragraph()
    p_tab_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tc = p_tab_cap.add_run(f'{tab_label} {caption}')
    set_run_font(run_tc, '宋体', size_pt=10.5, bold=True)

    table = doc.add_table(rows=len(body_rows) + 1, cols=ncols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_three_line_table(table)

    # 表头行设为重复标题行（跨页自动出现）
    tblHeader_el = OxmlElement('w:tblHeader')
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(tblHeader_el)

    for j, cell_children in enumerate(head_cells):
        cell = table.rows[0].cells[j]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _set_para_mark_size(p, 9)
        walk_inline(p, cell_children, base_size=9)

    for i, row_data in enumerate(body_rows):
        for j, cell_children in enumerate(row_data):
            if j < ncols:
                cell = table.rows[i + 1].cells[j]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                _set_para_mark_size(p, 9)
                walk_inline(p, cell_children, base_size=9)

    add_empty_para(doc)


def _render_display_math(doc, math_text, chapter_path, eq_counter, has_chapter):
    """渲染行间公式：上下各空一行、OMLL 居中、编号右对齐

    参数：
        doc: Document 对象
        math_text: LaTeX 公式文本（不含 $$ 定界符）
        chapter_path: 章节路径 list
        eq_counter: 公式计数 dict
        has_chapter: 是否有章标题
    """
    try:
        omml = latex_to_omml(math_text, display=True)
    except Exception:
        omml = None

    if omml is not None and omml.tag == f'{{{NSM}}}oMathPara':
        # oMathPara 可能来自 align 多行环境，或 latex_to_omml(display=True) 对单行公式的包装
        omath_children = [c for c in list(omml) if c.tag == f'{{{NSM}}}oMath']
        if len(omath_children) > 1:
            # align 等多行公式：oMath 以 <w:br/> 分隔
            add_empty_para(doc)
            p_align = doc.add_paragraph()
            p_align.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for idx, line_omml in enumerate(omath_children):
                if idx > 0:
                    br_run = p_align.add_run()
                    br = OxmlElement('w:br')
                    br_run._r.append(br)
                p_align._element.append(line_omml)
            add_empty_para(doc)
            return
        # 单行公式：提取 oMath，走下方编号路径
        omml = omath_children[0] if omath_children else omml

    add_empty_para(doc)
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    center_tab = OxmlElement('w:tab')
    center_tab.set(qn('w:val'), 'center')
    center_tab.set(qn('w:pos'), '4536')
    tabs.append(center_tab)
    right_tab = OxmlElement('w:tab')
    right_tab.set(qn('w:val'), 'right')
    right_tab.set(qn('w:pos'), '9072')
    tabs.append(right_tab)
    pPr.append(tabs)
    run_t1 = p.add_run()
    tab1 = OxmlElement('w:tab')
    run_t1._r.append(tab1)
    if omml is not None:
        p._element.append(omml)
    else:
        run_fb = p.add_run('[公式解析失败]')
        set_run_font(run_fb, '宋体', en_font='Times New Roman', size_pt=10.5)
    run_t2 = p.add_run()
    tab2 = OxmlElement('w:tab')
    run_t2._r.append(tab2)
    ch_num = chapter_path[0]

    eq_key = tuple(chapter_path[:1]) if has_chapter else None
    if eq_key:
        eq_counter[eq_key] = eq_counter.get(eq_key, 0) + 1
        eq_num = eq_counter[eq_key]
    else:
        eq_counter[None] = eq_counter.get(None, 0) + 1
        eq_num = eq_counter[None]

    run_lp = p.add_run('(')
    set_run_font(run_lp, '宋体', en_font='宋体', size_pt=10.5)
    run_lnum = p.add_run(f'{ch_num}-{eq_num}')
    set_run_font(run_lnum, '宋体', en_font='Times New Roman', size_pt=10.5)
    run_rp = p.add_run(')')
    set_run_font(run_rp, '宋体', en_font='宋体', size_pt=10.5)
    add_empty_para(doc)


def generate_docx(ast, output_path, title_text=None):
    """遍历 mistune AST 生成 docx 文件

    参数：
        ast: mistune AST 节点列表
        output_path: 输出 docx 文件路径
        title_text: 文档题目（用于页眉）
    """
    doc = Document()
    override_builtin_heading_styles(doc)

    # --- 默认样式 ---
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # --- 节/页边距设置 ---
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    header = section.header
    header.is_linked_to_previous = False
    hp_left = header.paragraphs[0]
    hp_left.paragraph_format.space_after = Pt(0)
    run_l = hp_left.add_run('xxxxx')
    set_run_font(run_l, '黑体', size_pt=9)
    hp_right = header.add_paragraph()
    hp_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp_right.paragraph_format.space_before = Pt(0)
    run_r = hp_right.add_run(title_text or '未命名文档')
    set_run_font(run_r, '黑体', size_pt=9)

    # --- 页脚/页码设置 ---
    section.footer_distance = Cm(1)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run('第')
    set_run_font(r1, '宋体', size_pt=10.5)
    r_page = fp.add_run()
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
    r_page._r.append(fc1)
    it1 = OxmlElement('w:instrText'); it1.set(qn('xml:space'), 'preserve'); it1.text = ' PAGE '
    r_page._r.append(it1)
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end')
    r_page._r.append(fc2)
    r2 = fp.add_run('页  共')
    set_run_font(r2, '宋体', size_pt=10.5)
    r_total = fp.add_run()
    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'begin')
    r_total._r.append(fc3)
    it2 = OxmlElement('w:instrText'); it2.set(qn('xml:space'), 'preserve'); it2.text = ' NUMPAGES '
    r_total._r.append(it2)
    fc4 = OxmlElement('w:fldChar'); fc4.set(qn('w:fldCharType'), 'end')
    r_total._r.append(fc4)
    r3 = fp.add_run('页')
    set_run_font(r3, '宋体', size_pt=10.5)

    # --- 列表编号定义 ---
    used_ids = _ensure_list_numbering_defs(doc)

    # --- 计数器 ---
    chapter_path = [1]  # 一级标题序号
    has_chapter = False
    fig_counter = {}
    tab_counter = {}
    eq_counter = {}

    heading_count = 0  # 跟踪标题，第一个 # 是题目
    pending_table_caption = None  # "表 xxx" 段落缓冲

    # --- AST Walker ---
    ast_len = len(ast)
    idx = 0
    while idx < ast_len:
        node = ast[idx]
        t = node['type']

        if t == 'heading':
            heading_count += 1
            level = node.get('attrs', {}).get('level', 1)
            text = _extract_ast_text(node.get('children', []))

            if level == 1 and heading_count == 1:
                # 题目：三号黑体居中
                add_empty_para(doc)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                set_run_font(run, '黑体', size_pt=16)
                add_empty_para(doc)
            else:
                if level == 1:
                    has_chapter = True
                    chapter_path[0] += 1

                if level == 1:
                    add_empty_para(doc)
                    p = doc.add_paragraph()
                    add_page_break_before(p)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_outline_level(p, 1)
                    run = p.add_run(text)
                    set_run_font(run, '黑体', size_pt=16)
                    add_empty_para(doc)
                elif level == 2:
                    p = doc.add_paragraph()
                    set_outline_level(p, 2)
                    # 顶格：段落级显式清零首行/左侧缩进（twips + chars 双维度）
                    ind = OxmlElement('w:ind')
                    for attr in ('w:left', 'w:leftChars', 'w:firstLine',
                                 'w:firstLineChars', 'w:hanging', 'w:hangingChars'):
                        ind.set(qn(attr), '0')
                    p._element.get_or_add_pPr().insert(0, ind)
                    run = p.add_run(text)
                    set_run_font(run, '黑体', size_pt=14, bold=False)
                elif level == 3:
                    p = doc.add_paragraph()
                    set_outline_level(p, 3)
                    set_first_line_indent_chars(p, 2)
                    run = p.add_run(text)
                    set_run_font(run, '宋体', size_pt=12, bold=False)
                else:
                    p = doc.add_paragraph()
                    set_outline_level(p, level)
                    set_first_line_indent_chars(p, 2)
                    p.paragraph_format.line_spacing = 1.3
                    run = p.add_run(text)
                    set_run_font(run, '宋体', size_pt=10.5, bold=False)

        elif t == 'paragraph':
            children = node.get('children', [])
            # 检测单图段落（仅含一个 image 子节点）：作为独立图片渲染
            if len(children) == 1 and children[0]['type'] == 'image':
                img_node = children[0]
                url = img_node.get('attrs', {}).get('url', img_node.get('src', ''))
                alt = img_node.get('attrs', {}).get('alt', img_node.get('alt', ''))
                img_path = download_image(url)
                if img_path is None:
                    img_path = add_placeholder_image(doc, alt or url)

                w, h = calc_image_size(img_path)

                add_empty_para(doc)
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(img_path, width=w)

                if not alt:
                    alt = os.path.splitext(os.path.basename(url))[0]

                fig_key = tuple(chapter_path[:1]) if has_chapter else None
                if fig_key:
                    fig_counter[fig_key] = fig_counter.get(fig_key, 0) + 1
                    fig_num = fig_counter[fig_key]
                else:
                    fig_counter[None] = fig_counter.get(None, 0) + 1
                    fig_num = fig_counter[None]

                fig_label = f"图{chapter_path[0]}-{fig_num}" if has_chapter else f"图{fig_num}"

                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_cap = p_cap.add_run(f'{fig_label} {alt}')
                set_run_font(run_cap, '宋体', size_pt=9, bold=True)
                add_empty_para(doc)

                try:
                    os.unlink(img_path)
                except Exception:
                    pass
            else:
                # 检测 "表 xxx" 段落：若下一个节点是 table，作为表题
                para_text = _extract_ast_text(children)
                # 检查下一个非空行节点是否为 table（跳过 blank_line）
                next_idx = idx + 1
                while next_idx < ast_len and ast[next_idx]['type'] == 'blank_line':
                    next_idx += 1
                if para_text.startswith('表') and next_idx < ast_len and ast[next_idx]['type'] == 'table':
                    # 缓冲为表题，去除"表 "前缀
                    pending_table_caption = para_text[1:].strip() if len(para_text) > 1 else para_text
                else:
                    p = doc.add_paragraph()
                    set_first_line_indent_chars(p, 2)
                    p.paragraph_format.line_spacing = 1.3
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    walk_inline(p, children)

        elif t == 'blank_line':
            pass  # 不打断表题缓冲

        elif t == 'thematic_break':
            p = doc.add_paragraph()
            run = p.add_run('')
            set_run_font(run, '宋体', size_pt=10.5)
            add_page_break_before(p)
            pending_table_caption = None

        elif t == 'list':
            _render_list(doc, node, used_ids=used_ids)
            pending_table_caption = None

        elif t == 'block_math':
            _render_display_math(doc, node['raw'], chapter_path, eq_counter, has_chapter)
            pending_table_caption = None

        elif t == 'block_code':
            add_empty_para(doc)
            p = doc.add_paragraph()
            set_first_line_indent_chars(p, 0)
            pPr = p._element.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'D9D9D9')
            pPr.append(shd)
            run = p.add_run(node.get('raw', ''))
            set_run_font(run, '宋体', en_font='Times New Roman', size_pt=10.5)
            add_empty_para(doc)
            pending_table_caption = None

        elif t == 'table':
            caption = pending_table_caption or ''
            pending_table_caption = None
            _render_table(doc, node, tab_counter, chapter_path, has_chapter, caption)

        idx += 1

    doc.save(output_path)
    return output_path


# ============================================================
# 主入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description='mddoc - Markdown 转学术格式 DOCX',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python md2docx.py paper.md                    # 转换文件，输出到同目录
  python md2docx.py paper.md -o output.docx     # 指定输出路径
  python md2docx.py --text "# 标题\\n\\n正文内容" -o out.docx  # 直接转换文本
        """.strip())
    parser.add_argument('input', nargs='?', help='输入 Markdown 文件路径')
    parser.add_argument('-o', '--output', help='输出 DOCX 文件路径')
    parser.add_argument('--text', help='直接传入 Markdown 文本（不需要文件）')
    args = parser.parse_args()

    if not args.input and not args.text:
        parser.print_help()
        sys.exit(1)

    # 读取内容
    if args.text:
        md_content = args.text
        source_dir = os.getcwd()
    else:
        with open(args.input, 'r', encoding='utf-8') as f:
            md_content = f.read()
        source_dir = os.path.dirname(os.path.abspath(args.input))

    # 解析
    nodes = parse_markdown(md_content)

    # 提取题目（第一个 heading level=1）
    title_node = None
    for n in nodes:
        if n['type'] == 'heading' and n.get('attrs', {}).get('level') == 1:
            title_node = n
            break
    title_text = _extract_ast_text(title_node['children']) if title_node else None

    # 确定输出路径
    if args.output:
        output_path = args.output
    elif title_text:
        safe_name = re.sub(r'[\\/*?:"<>|]', '', title_text)
        output_path = os.path.join(source_dir, f'{safe_name}.docx')
    elif args.text:
        output_path = os.path.join(source_dir, '未命名文档.docx')
    else:
        stem = os.path.splitext(os.path.basename(args.input))[0]
        output_path = os.path.join(source_dir, f'{stem}.docx')

    # 生成
    generate_docx(nodes, output_path, title_text)
    print(f'已生成: {output_path}')


if __name__ == '__main__':
    main()
